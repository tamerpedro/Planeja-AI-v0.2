import hashlib
import io
import json
import math
import re
from datetime import datetime
from importlib.metadata import PackageNotFoundError, version
from pathlib import Path
from zoneinfo import ZoneInfo

import streamlit as st
from openai import APIStatusError, AuthenticationError, BadRequestError, OpenAI, RateLimitError


st.set_page_config(page_title="Planeja IA - DOD v0.4", page_icon="D", layout="wide")
st.title("Planeja IA - Minuta de DOD")
st.caption("Assistente para estruturar minutas de Documento de Oficializacao de Demanda da Dataprev.")

MOTIVOS = [
    "Termino do contrato vigente",
    "Obsolescencia da tecnologia",
    "Adequacao as boas praticas de mercado",
    "Exigencia legal ou normativa",
    "Solicitacao dos clientes",
    "Outros",
]

EIXOS = [
    "PRODUTOS E SOLUCOES",
    "GESTAO TECNOLOGICA, ANALYTICS E IA",
    "SEGURANCA E RESILIENCIA",
    "GESTAO E GOVERNANCA",
    "PESSOAS",
]

DIRETRIZES_PDTIC = [
    "D1. Conceber solucoes buscando promover a independencia tecnologica em relacao a fornecedores.",
    "D2. Atualizar, disponibilizar e adequar a infraestrutura de TI as necessidades do negocio.",
    "D3. Entregar solucoes tecnologicas seguras, usando melhores praticas de mercado.",
    "D4. Entregar solucoes com padroes e protocolos amplamente utilizados no mercado.",
    "D5. Entregar solucoes tecnologicas modernas e inovadoras, baseadas em dados.",
    "D6. Realizar, quando oportuno, parcerias estrategicas com fornecedores de tecnologia.",
    "D7. Adquirir solucoes tecnologicas de uso corporativo seguindo padroes de mercado.",
    "D8. Prospectar solucoes inovadoras considerando oportunidades, riscos tecnologicos e seguranca.",
    "D9. Promover fortalecimento do corpo funcional.",
    "D10. Priorizar o uso da automacao nos processos.",
    "D11. Promover ambiente digital de trabalho colaborativo, eficiente e seguro.",
    "D12. Promover contratacoes de solucoes tecnologicas a longo prazo.",
    "D13. Contratar solucoes prevendo capacidade de crescimento futuro.",
    "D14. Fomentar cultura orientada a dados em todos os niveis da organizacao.",
]

PILARES_PDTIC = [
    "Manutencao e Expansao",
    "Aprimoramento da Gestao da Continuidade de Negocios",
    "Consolidacao da Gestao de Servicos de TI",
    "Modernizacao",
    "Adocao de novas tecnologias de desenvolvimento",
    "Consolidacao da Modernizacao Continua",
    "Implementacao de Hiperautomacao e integradores",
    "Inovacao continua adotando IA para modelos avancados",
    "Evolucao e Inovacao",
    "Implementacao de Solucoes em Multinuvem de Governo",
    "Ampliacao dos processos de Inteligencia Analitica",
    "Consolidacao da Infraestrutura Publica Digital",
    "Fortalecimento da Infraestrutura Nacional de Dados",
]

SERVICOS = ["Orientacao Tecnica", "Capacitacao Tecnica", "Suporte Tecnico"]

TIPOS_ANEXO = [
    "Modelo oficial de DOD",
    "Normativo interno aplicavel",
    "Legislacao aplicavel",
    "DOD anterior semelhante",
    "Especificacao tecnica previa do objeto",
    "Documento de exemplo",
    "Outro contexto de apoio",
]

CATEGORIAS_BIBLIOTECA = {
    "Modelo oficial de DOD": ("modelos", 1),
    "Normativo interno Dataprev": ("normativos_dataprev", 2),
    "Legislacao aplicavel": ("legislacao", 3),
    "DOD anterior semelhante": ("dods_anteriores", 4),
    "Referencia tecnica": ("referencias_tecnicas", 5),
}

PASTA_PARA_CATEGORIA = {
    pasta: tipo
    for tipo, (pasta, _prioridade) in CATEGORIAS_BIBLIOTECA.items()
}

BASE_INSTITUCIONAL = Path("base_institucional")
INDICE_INSTITUCIONAL = BASE_INSTITUCIONAL / "indice_contexto.json"
ARQUIVOS_IGNORADOS_BIBLIOTECA = {".gitkeep", "indice_contexto.json"}

ESTRUTURA_DOD = """
Documento de Oficializacao de Demanda
Participantes: Elaboracao e Aprovacao
Historico de revisoes do documento
1. IDENTIFICACAO DA AREA DEMANDANTE DA SOLUCAO
2. IDENTIFICACAO DA DEMANDA
2.1. CONTEXTO DE NEGOCIO
3. CONTEXTO DA DEMANDA
3.1. SITUACAO ATUAL
3.2. ESCOPO DA DEMANDA
3.3. MOTIVACAO DA DEMANDA
3.3.1. Assinalar motivacao
3.3.2. Riscos envolvidos caso a contratacao nao seja realizada
3.3.3. Resultados a serem alcancados
3.4. DATA PREVISTA PARA DISPONIBILIZACAO DA DEMANDA
3.5. FORNECEDOR(ES), SE HOUVER
3.6. DESCRICAO DOS OBJETOS E QUANTIDADES ENVOLVIDAS
3.6.1. Para contratos existentes
3.6.2. Para nova contratacao
3.7. SERVICOS ASSOCIADOS A DEMANDA
3.7.1. Servicos selecionados
3.7.2. Condicoes minimas obrigatorias
4. AREAS E PAPEIS ENVOLVIDOS
4.1. AREAS INTERNAS (DATAPREV)
4.2. CLIENTES EXTERNOS QUE FARAO USO DA SOLUCAO/SOFTWARE
5. INFORMACOES ADICIONAIS
6. NOTAS
7. ANEXOS
"""


def carregar_contexto() -> str:
    caminho = Path("contexto.md")
    return caminho.read_text(encoding="utf-8") if caminho.exists() else ""


def obter_secret_texto(nome: str, obrigatorio: bool = False) -> str:
    valor = str(st.secrets.get(nome, "") or "").strip()
    if obrigatorio and not valor:
        st.error(f"{nome} nao configurada nos Secrets do Streamlit Cloud.")
        st.stop()
    return valor


def obter_cliente_openai() -> OpenAI:
    return OpenAI(api_key=obter_secret_texto("OPENAI_API_KEY", obrigatorio=True))


def obter_modelo() -> str:
    return obter_secret_texto("OPENAI_MODEL") or "gpt-5.5"


def obter_modelo_embedding() -> str:
    return obter_secret_texto("OPENAI_EMBEDDING_MODEL") or "text-embedding-3-small"


def obter_versao_openai_sdk() -> str:
    try:
        return version("openai")
    except PackageNotFoundError:
        return "nao identificada"


def obter_timestamp_brasilia() -> str:
    return datetime.now(ZoneInfo("America/Sao_Paulo")).isoformat()


def mostrar_erro_openai(titulo: str, erro: Exception) -> None:
    request_id = getattr(erro, "request_id", "") or ""
    response = getattr(erro, "response", None)
    if not request_id and response is not None:
        request_id = response.headers.get("x-request-id", "")

    st.error(titulo)
    st.write(f"**Timestamp:** `{obter_timestamp_brasilia()}`")
    st.write("**Timezone:** `America/Sao_Paulo`")
    st.write(f"**OpenAI Python SDK:** `{obter_versao_openai_sdk()}`")
    st.write("**Endpoint:** `Responses API - client.responses.create`")
    if getattr(erro, "status_code", None):
        st.write(f"**Status code:** `{erro.status_code}`")
    if request_id:
        st.write(f"**x-request-id:** `{request_id}`")
    st.write("**Detalhes do erro:**")
    st.code(str(erro))
    st.stop()


def campo_texto(label: str, key: str, height: int = 110, placeholder: str = "") -> str:
    return st.text_area(label, key=key, height=height, placeholder=placeholder)


def limitar_texto(texto: str, limite: int = 12000) -> str:
    texto = texto.strip()
    if len(texto) <= limite:
        return texto
    return texto[:limite] + "\n\n[Texto truncado pelo app para preservar limite de contexto.]"


def preparar_biblioteca_institucional() -> None:
    BASE_INSTITUCIONAL.mkdir(exist_ok=True)
    for pasta, _prioridade in CATEGORIAS_BIBLIOTECA.values():
        (BASE_INSTITUCIONAL / pasta).mkdir(exist_ok=True)


def normalizar_nome_arquivo(nome: str) -> str:
    nome_limpo = Path(nome).name
    nome_limpo = re.sub(r"[^A-Za-z0-9._ -]+", "_", nome_limpo).strip()
    return nome_limpo or "documento"


def obter_nome_arquivo(arquivo) -> str:
    return getattr(arquivo, "name", Path(str(arquivo)).name)


def obter_bytes_arquivo(arquivo) -> bytes:
    if isinstance(arquivo, Path):
        return arquivo.read_bytes()

    if isinstance(arquivo, str):
        return Path(arquivo).read_bytes()

    return arquivo.getvalue()


def carregar_indice_institucional() -> list[dict]:
    if not INDICE_INSTITUCIONAL.exists():
        return []

    try:
        return json.loads(INDICE_INSTITUCIONAL.read_text(encoding="utf-8"))
    except Exception:
        return []


def salvar_indice_institucional(indice: list[dict]) -> None:
    preparar_biblioteca_institucional()
    INDICE_INSTITUCIONAL.write_text(
        json.dumps(indice, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def dividir_em_trechos(texto: str, tamanho: int = 1800, sobreposicao: int = 250) -> list[str]:
    texto = re.sub(r"\n{3,}", "\n\n", texto).strip()
    if not texto:
        return []

    trechos = []
    inicio = 0

    while inicio < len(texto):
        fim = min(inicio + tamanho, len(texto))
        corte = texto.rfind("\n\n", inicio, fim)

        if corte > inicio + 500:
            fim = corte

        trecho = texto[inicio:fim].strip()
        if trecho:
            trechos.append(trecho)

        if fim >= len(texto):
            break

        inicio = max(0, fim - sobreposicao)

    return trechos


def obter_embeddings(client: OpenAI, textos: list[str]) -> list[list[float]]:
    if not textos:
        return []

    resposta = client.embeddings.create(
        model=obter_modelo_embedding(),
        input=textos,
    )
    return [item.embedding for item in resposta.data]


def similaridade_cosseno(a: list[float], b: list[float]) -> float:
    produto = sum(x * y for x, y in zip(a, b))
    norma_a = math.sqrt(sum(x * x for x in a))
    norma_b = math.sqrt(sum(y * y for y in b))

    if not norma_a or not norma_b:
        return 0.0

    return produto / (norma_a * norma_b)


def montar_consulta_institucional(dados: dict, anexos: list[dict[str, str]]) -> str:
    partes = [
        dados.get("titulo_dod", ""),
        dados.get("nome_projeto", ""),
        dados.get("unidade_demandante", ""),
        dados.get("contexto_negocio", ""),
        dados.get("situacao_atual", ""),
        dados.get("escopo_demanda", ""),
        dados.get("riscos", ""),
        dados.get("resultados", ""),
        dados.get("objetos_quantidades", ""),
        dados.get("justificativa_quantidades", ""),
        " ".join(dados.get("motivacoes", [])),
        " ".join(dados.get("diretrizes_pdtic", [])),
        " ".join(dados.get("pilares_pdtic", [])),
    ]

    for anexo in anexos:
        partes.append(anexo.get("descricao", ""))
        partes.append(anexo.get("conteudo_extraido", "")[:1000])

    return "\n".join(str(parte) for parte in partes if str(parte).strip())


def buscar_contexto_institucional(client: OpenAI, dados: dict, anexos: list[dict[str, str]], limite: int = 8) -> list[dict]:
    indice = carregar_indice_institucional()
    if not indice:
        return []

    consulta = montar_consulta_institucional(dados, anexos)
    if not consulta.strip():
        return []

    embedding_consulta = obter_embeddings(client, [consulta[:6000]])[0]
    resultados = []

    for item in indice:
        embedding = item.get("embedding")
        if not embedding:
            continue

        score = similaridade_cosseno(embedding_consulta, embedding)
        prioridade = int(item.get("prioridade", 99))
        resultados.append((score, prioridade, item))

    resultados.sort(key=lambda item: (item[0], -item[1]), reverse=True)

    selecionados = []
    for score, _prioridade, item in resultados[:limite]:
        copia = dict(item)
        copia["score"] = round(score, 4)
        copia.pop("embedding", None)
        selecionados.append(copia)

    return selecionados


def formatar_contexto_institucional(trechos: list[dict]) -> str:
    if not trechos:
        return "Nenhum trecho institucional recuperado para esta demanda."

    blocos = []
    for indice, item in enumerate(trechos, start=1):
        blocos.append(
            "\n".join(
                [
                    f"[{indice}] Fonte: {item.get('arquivo', '<<sem arquivo>>')}",
                    f"Tipo: {item.get('tipo_documental', '<<sem tipo>>')}",
                    f"Assunto: {item.get('assunto') or '<<nao informado>>'}",
                    f"Area relacionada: {item.get('area_relacionada') or '<<nao informada>>'}",
                    f"Versao/Data: {item.get('versao_data') or '<<nao informada>>'}",
                    f"Sensibilidade: {item.get('sensibilidade') or '<<nao informada>>'}",
                    f"Score: {item.get('score')}",
                    "Trecho:",
                    item.get("texto", ""),
                ]
            )
        )

    return "\n\n---\n\n".join(blocos)


def indexar_documentos_institucionais(
    client: OpenAI,
    arquivos,
    tipo_documental: str,
    assunto: str,
    area_relacionada: str,
    versao_data: str,
    sensibilidade: str,
) -> int:
    preparar_biblioteca_institucional()
    pasta, prioridade = CATEGORIAS_BIBLIOTECA[tipo_documental]
    indice = carregar_indice_institucional()
    novos_itens = []

    for arquivo in arquivos:
        conteudo_bytes = obter_bytes_arquivo(arquivo)
        nome = normalizar_nome_arquivo(obter_nome_arquivo(arquivo))
        hash_documento = hashlib.sha256(conteudo_bytes).hexdigest()
        caminho = BASE_INSTITUCIONAL / pasta / nome

        if not isinstance(arquivo, Path) or caminho.resolve() != arquivo.resolve():
            caminho.write_bytes(conteudo_bytes)

        texto = extrair_texto_anexo(arquivo)
        trechos = dividir_em_trechos(texto)
        embeddings = obter_embeddings(client, trechos)

        indice = [
            item for item in indice
            if item.get("hash_documento") != hash_documento
        ]

        for numero, (trecho, embedding) in enumerate(zip(trechos, embeddings), start=1):
            novos_itens.append(
                {
                    "id": f"{hash_documento}:{numero}",
                    "hash_documento": hash_documento,
                    "arquivo": nome,
                    "caminho": str(caminho.as_posix()),
                    "tipo_documental": tipo_documental,
                    "pasta": pasta,
                    "prioridade": prioridade,
                    "assunto": assunto.strip(),
                    "area_relacionada": area_relacionada.strip(),
                    "versao_data": versao_data.strip(),
                    "sensibilidade": sensibilidade,
                    "trecho": numero,
                    "texto": trecho,
                    "embedding_model": obter_modelo_embedding(),
                    "embedding": embedding,
                }
            )

    salvar_indice_institucional(indice + novos_itens)
    return len(novos_itens)


def listar_documentos_do_repositorio() -> list[tuple[Path, str]]:
    preparar_biblioteca_institucional()
    documentos = []

    for pasta, tipo_documental in PASTA_PARA_CATEGORIA.items():
        diretorio = BASE_INSTITUCIONAL / pasta
        if not diretorio.exists():
            continue

        for caminho in diretorio.rglob("*"):
            if not caminho.is_file():
                continue

            if caminho.name in ARQUIVOS_IGNORADOS_BIBLIOTECA:
                continue

            documentos.append((caminho, tipo_documental))

    return documentos


def reindexar_biblioteca_do_repositorio(client: OpenAI) -> tuple[int, int]:
    documentos = listar_documentos_do_repositorio()
    indice = []
    total_trechos = 0

    salvar_indice_institucional([])

    for caminho, tipo_documental in documentos:
        assunto = caminho.stem.replace("_", " ").replace("-", " ").strip()
        total_trechos += indexar_documentos_institucionais(
            client,
            [caminho],
            tipo_documental,
            assunto,
            "",
            "",
            "Uso interno autorizado",
        )
        indice = carregar_indice_institucional()

    return len(documentos), len(indice) if indice else total_trechos


def extrair_texto_anexo(arquivo) -> str:
    nome = obter_nome_arquivo(arquivo).lower()
    try:
        if nome.endswith((".xlsx", ".xls")):
            planilhas = []
            folhas = st.session_state.get("_limite_folhas_planilha", 10)
            linhas = st.session_state.get("_limite_linhas_planilha", 80)
            workbook = __import__("pandas").read_excel(
                io.BytesIO(obter_bytes_arquivo(arquivo)),
                sheet_name=None,
                nrows=linhas,
            )
            for nome_folha, dados in list(workbook.items())[:folhas]:
                planilhas.append(f"### Planilha: {nome_folha}\n{dados.to_markdown(index=False)}")
            return "\n\n".join(planilhas)

        if nome.endswith(".csv"):
            dados = __import__("pandas").read_csv(io.BytesIO(obter_bytes_arquivo(arquivo)), nrows=300)
            return dados.to_markdown(index=False)

        if nome.endswith(".pdf"):
            from pypdf import PdfReader

            leitor = PdfReader(io.BytesIO(obter_bytes_arquivo(arquivo)))
            return "\n\n".join(p.extract_text() or "" for p in leitor.pages)

        if nome.endswith(".docx"):
            from docx import Document

            documento = Document(io.BytesIO(obter_bytes_arquivo(arquivo)))
            paragrafos = [p.text for p in documento.paragraphs]
            tabelas = [
                " | ".join(c.text.strip() for c in linha.cells)
                for tabela in documento.tables
                for linha in tabela.rows
            ]
            return "\n".join(paragrafos + tabelas)

        return obter_bytes_arquivo(arquivo).decode("utf-8")
    except UnicodeDecodeError:
        return "[Arquivo binario sem extracao automatica. Descreva o conteudo no campo de observacoes.]"
    except Exception as erro:
        return f"[Nao foi possivel extrair o texto automaticamente: {erro}]"


def montar_anexos_contexto(arquivos) -> list[dict[str, str]]:
    anexos = []
    for indice, arquivo in enumerate(arquivos, start=1):
        with st.expander(f"Classificar anexo {indice}: {arquivo.name}", expanded=True):
            tipo = st.selectbox("Tipo de anexo", TIPOS_ANEXO, key=f"tipo_{indice}_{arquivo.name}")
            descricao = st.text_area(
                "Como este anexo deve orientar a minuta?",
                key=f"desc_{indice}_{arquivo.name}",
                height=90,
                placeholder="Ex.: DOD anterior para demanda semelhante; usar apenas como referencia.",
            )
            conteudo = limitar_texto(extrair_texto_anexo(arquivo))
            st.caption(f"Texto extraido: {len(conteudo)} caracteres")
        anexos.append({"arquivo": arquivo.name, "tipo": tipo, "descricao": descricao, "conteudo_extraido": conteudo})
    return anexos


def montar_prompt(dados: dict, contexto: str, anexos: list[dict[str, str]], contexto_institucional: str) -> str:
    return f"""
Voce e um assistente especializado na elaboracao de minutas de Documento de Oficializacao de Demanda (DOD) para a Dataprev.

Objetivo: gerar uma minuta em texto, formal e semi estruturada, seguindo o mais fielmente possivel a estrutura do modelo oficial de DOD da Dataprev.

Ordem de prioridade dos contextos:
1. Modelo oficial de DOD da Dataprev e estrutura obrigatoria informada abaixo.
2. Normativos internos aplicaveis fornecidos no contexto ou anexos.
3. Legislacao aplicavel fornecida no contexto ou anexos.
4. DODs anteriores semelhantes, apenas como referencia de estilo, estrutura e nivel de detalhe.
5. Especificacoes tecnicas e demais anexos fornecidos pelo usuario.
6. Biblioteca institucional recuperada por relevancia, como contexto secundario e rastreavel.

Regras:
- Nao invente normas, valores, numeros de processo, areas, matriculas, fornecedores, prazos, quantitativos ou fatos nao fornecidos.
- Quando faltar informacao obrigatoria, use <<preencher>>.
- Use "Nao se aplica" somente quando os dados indicarem isso.
- Nao antecipe estrategia de contratacao, salvo se o usuario declarar expressamente.
- Remova orientacoes internas do modelo. A saida deve ser a minuta, nao instrucoes.
- Mantenha titulos e numeracao do DOD.
- Use "(X)" nas opcoes selecionadas e "( )" nas opcoes nao selecionadas.
- No item 2, apresente todos os eixos, diretrizes e pilares informados pelo usuario, preferencialmente em tabela.
- No item 3.2, preserve tabelas de lote, unidade, endereco, item, quantidade e garantia quando esses dados existirem no input ou anexos.
- No item 3.3.1, mostre todas as opcoes oficiais de motivacao, marcando apenas as selecionadas.
- No item 3.3.2, quando houver varios riscos, organize em categorias como riscos operacionais, tecnologicos, de gestao e institucionais.
- No item 3.3.3, desenvolva cada resultado esperado em paragrafo curto, e nao apenas em lista, quando houver informacao suficiente.
- No item 3.6, se houver quantitativo por unidade ou planilha anexa, consolide uma tabela por unidade e uma tabela de quantitativo total.
- No item 7, liste todos os anexos usados e descreva como cada um foi considerado.
- Use a biblioteca institucional apenas para orientar estrutura, linguagem, referencias e padroes. Nao copie fatos de DODs anteriores como se fossem dados da demanda atual.
- Se a biblioteca institucional trouxer legislacao ou normativos, use-os como orientacao. Nao invente enquadramento juridico, modalidade, dispensa, inexigibilidade, aditivo ou prorrogacao sem fonte explicita.
- Considere internamente as fontes recuperadas, mas nao polua a minuta final com citacoes longas.

Estrutura oficial:
{ESTRUTURA_DOD}

Contexto do repositorio:
{contexto or "Nenhum contexto adicional no repositorio."}

Dados estruturados:
{json.dumps(dados, ensure_ascii=False, indent=2)}

Anexos:
{json.dumps(anexos, ensure_ascii=False, indent=2)}

Biblioteca institucional recuperada:
{contexto_institucional}

Saida esperada: gere apenas a minuta do DOD em Markdown, com tabelas quando fizer sentido.
"""


with st.sidebar:
    contexto_repo = carregar_contexto()
    preparar_biblioteca_institucional()
    st.header("Configuracao")
    st.write("**OPENAI_API_KEY:**", "configurada" if st.secrets.get("OPENAI_API_KEY") else "nao configurada")
    st.write("**OPENAI_MODEL:**", obter_modelo())
    st.write("**OPENAI_EMBEDDING_MODEL:**", obter_modelo_embedding())
    st.write("**OpenAI Python SDK:**", obter_versao_openai_sdk())
    st.write("**OPENAI_ORG_ID:** nao utilizado")
    st.session_state["_limite_folhas_planilha"] = st.number_input(
        "Folhas por planilha",
        min_value=1,
        max_value=50,
        value=10,
    )
    st.session_state["_limite_linhas_planilha"] = st.number_input(
        "Linhas por folha",
        min_value=20,
        max_value=1000,
        value=80,
        step=20,
    )
    with st.expander("Contexto do repositorio"):
        st.markdown(contexto_repo or "Nenhum contexto encontrado.")

    with st.expander("Biblioteca institucional"):
        indice_atual = carregar_indice_institucional()
        documentos_indexados = {item.get("hash_documento") for item in indice_atual}
        st.write(f"**Documentos indexados:** {len(documentos_indexados)}")
        st.write(f"**Trechos indexados:** {len(indice_atual)}")
        documentos_repositorio = listar_documentos_do_repositorio()
        st.write(f"**Documentos no repositorio:** {len(documentos_repositorio)}")

        if st.button("Reindexar biblioteca do repositorio", key="reindexar_biblioteca_repositorio"):
            with st.spinner("Reindexando documentos do repositorio..."):
                try:
                    total_documentos, total_trechos = reindexar_biblioteca_do_repositorio(
                        obter_cliente_openai()
                    )
                    st.success(
                        f"{total_documentos} documentos processados; {total_trechos} trechos indexados."
                    )
                except Exception as erro:
                    st.error("Nao foi possivel reindexar a biblioteca do repositorio.")
                    st.code(str(erro))

        arquivos_biblioteca = st.file_uploader(
            "Adicionar documentos autorizados",
            accept_multiple_files=True,
            key="upload_biblioteca_institucional",
            help="Use apenas documentos autorizados para compor a base institucional.",
        )
        tipo_biblioteca = st.selectbox(
            "Tipo documental",
            list(CATEGORIAS_BIBLIOTECA.keys()),
            key="tipo_biblioteca",
        )
        assunto_biblioteca = st.text_input("Assunto", key="assunto_biblioteca")
        area_biblioteca = st.text_input("Area relacionada", key="area_biblioteca")
        versao_biblioteca = st.text_input("Versao/Data", key="versao_biblioteca")
        sensibilidade_biblioteca = st.selectbox(
            "Nivel de sensibilidade",
            ["Uso interno autorizado", "Divulgacao opcional", "Publico", "Restrito"],
            key="sensibilidade_biblioteca",
        )
        autorizado_biblioteca = st.checkbox(
            "Confirmo que os documentos foram autorizados para uso nesta biblioteca",
            key="autorizado_biblioteca",
        )

        if st.button("Salvar e indexar documentos", key="indexar_biblioteca"):
            if not arquivos_biblioteca:
                st.warning("Selecione ao menos um documento.")
            elif not autorizado_biblioteca:
                st.warning("Confirme a autorizacao antes de indexar.")
            else:
                with st.spinner("Indexando biblioteca institucional..."):
                    try:
                        total_trechos = indexar_documentos_institucionais(
                            obter_cliente_openai(),
                            arquivos_biblioteca,
                            tipo_biblioteca,
                            assunto_biblioteca,
                            area_biblioteca,
                            versao_biblioteca,
                            sensibilidade_biblioteca,
                        )
                        st.success(f"{total_trechos} trechos indexados.")
                    except Exception as erro:
                        st.error("Nao foi possivel indexar os documentos.")
                        st.code(str(erro))

st.subheader("1. Identificacao")
c1, c2 = st.columns(2)
with c1:
    titulo_dod = st.text_input("Titulo do DOD")
    identificador = st.text_input("Identificador interno (opcional)")
    unidade_demandante = st.text_area("Unidade demandante", height=90, placeholder="Diretoria / Superintendencia / Departamento / Divisao / Setor")
    responsavel_demanda = st.text_input("Responsavel pela Unidade Demandante")
    matricula_responsavel = st.text_input("Matricula do responsavel")
with c2:
    elaboradores = st.text_area("Elaboracao: nomes, lotacoes e matriculas", height=120)
    aprovadores = st.text_area("Aprovacao: nome, cargo/area e matricula", height=120)
    pontos_focais = st.text_area("Pontos focais", height=100, placeholder="Nome completo, lotacao, cargo e matricula")

st.subheader("2. Identificacao da demanda")
c1, c2 = st.columns(2)
with c1:
    nome_projeto = st.text_input("Nome do projeto")
    eixos = st.multiselect("Eixos e objetivos estrategicos", EIXOS)
    diretrizes = st.multiselect("Diretrizes do PDTIC", DIRETRIZES_PDTIC)
with c2:
    pilares = st.multiselect("Pilares tecnologicos do PDTIC", PILARES_PDTIC)
    contexto_negocio = campo_texto("2.1 Contexto de negocio", "contexto_negocio", placeholder="Circunstancias, fatos de origem e beneficios esperados.")

st.subheader("3. Contexto da demanda")
situacao_atual = campo_texto("3.1 Situacao atual", "situacao_atual", placeholder="Como a necessidade e atendida hoje.")
escopo_demanda = campo_texto("3.2 Escopo da demanda", "escopo_demanda", 140, "Necessidades funcionais, limitacoes, clientes e ambiente de uso.")

c1, c2 = st.columns(2)
with c1:
    motivacoes = st.multiselect("3.3.1 Motivacao da demanda", MOTIVOS)
    motivacao_outros = st.text_input("Se marcou 'Outros', cite quais")
    riscos = campo_texto("3.3.2 Riscos se a contratacao nao for realizada", "riscos")
    resultados = campo_texto("3.3.3 Resultados a serem alcancados", "resultados")
with c2:
    data_disponibilizacao = st.text_input("3.4 Data prevista para disponibilizacao", placeholder="Ex.: A partir de dd/mm/aaaa, em virtude de...")
    fornecedores = campo_texto("3.5 Fornecedor(es), se houver", "fornecedores", placeholder="Fornecedor, contato, e-mail e telefone.")

st.subheader("3.6 Objetos e quantidades")
tipo_contratacao = st.radio("Tipo de demanda para dimensionamento", ["Nova contratacao", "Contrato existente", "Ainda nao definido"], horizontal=True)
objetos_quantidades = campo_texto("Descricao dos objetos e quantidades envolvidas", "objetos_quantidades", 130, "Itens, quantidades, acessos, usuarios, volume ou memoria de calculo.")
justificativa_quantidades = campo_texto("Justificativa das quantidades ou dados de dimensionamento", "justificativa_quantidades")

st.subheader("3.7 Servicos associados")
servicos = st.multiselect("Servicos que devem ser contemplados", SERVICOS)
c1, c2, c3 = st.columns(3)
with c1:
    orientacao_tecnica = campo_texto("Orientacao Tecnica", "orientacao_tecnica", placeholder="Horas e justificativa.")
with c2:
    capacitacao_tecnica = campo_texto("Capacitacao Tecnica", "capacitacao_tecnica", placeholder="Treinandos, areas e justificativa.")
with c3:
    suporte_tecnico = campo_texto("Suporte Tecnico", "suporte_tecnico", placeholder="Modalidade 24x7 ou 8x5 e justificativa.")

st.subheader("4. Areas e papeis envolvidos")
matriz_responsabilidades = st.text_area(
    "4.1 Areas internas e matriz de responsabilidades",
    height=150,
    placeholder="Gestor do Produto/Servico/Solucao: ...\nGestor Tecnico do Contrato: ...\nUsuarios: ...\nInstalacao, Operacao e Sustentacao: ...",
)
clientes_externos = campo_texto("4.2 Clientes externos beneficiados ou usuarios da solucao", "clientes_externos")

st.subheader("5. Informacoes adicionais, notas e anexos")
informacoes_adicionais = campo_texto("5. Informacoes adicionais", "informacoes_adicionais")
notas = campo_texto("6. Notas explicativas ou referencias", "notas")
arquivos = st.file_uploader(
    "7. Anexos de contexto",
    accept_multiple_files=True,
    help="Aceita qualquer extensao. O app extrai texto de PDF, DOCX, planilhas, CSV, TXT, MD e JSON; outros formatos entram como anexo descrito pelo usuario.",
)
anexos_contexto = montar_anexos_contexto(arquivos)

dados_dod = {
    "titulo_dod": titulo_dod,
    "identificador": identificador,
    "unidade_demandante": unidade_demandante,
    "responsavel_demanda": responsavel_demanda,
    "matricula_responsavel": matricula_responsavel,
    "elaboradores": elaboradores,
    "aprovadores": aprovadores,
    "pontos_focais": pontos_focais,
    "nome_projeto": nome_projeto,
    "eixos_estrategicos": eixos,
    "diretrizes_pdtic": diretrizes,
    "pilares_pdtic": pilares,
    "contexto_negocio": contexto_negocio,
    "situacao_atual": situacao_atual,
    "escopo_demanda": escopo_demanda,
    "motivacoes": motivacoes,
    "motivacao_outros": motivacao_outros,
    "riscos": riscos,
    "resultados": resultados,
    "data_disponibilizacao": data_disponibilizacao,
    "fornecedores": fornecedores,
    "tipo_contratacao": tipo_contratacao,
    "objetos_quantidades": objetos_quantidades,
    "justificativa_quantidades": justificativa_quantidades,
    "servicos_associados": servicos,
    "orientacao_tecnica": orientacao_tecnica,
    "capacitacao_tecnica": capacitacao_tecnica,
    "suporte_tecnico": suporte_tecnico,
    "matriz_responsabilidades": matriz_responsabilidades,
    "clientes_externos": clientes_externos,
    "informacoes_adicionais": informacoes_adicionais,
    "notas": notas,
}

campos_minimos = {
    "Titulo do DOD": titulo_dod,
    "Unidade demandante": unidade_demandante,
    "Responsavel pela demanda": responsavel_demanda,
    "Nome do projeto": nome_projeto,
    "Contexto de negocio": contexto_negocio,
    "Situacao atual": situacao_atual,
    "Escopo da demanda": escopo_demanda,
    "Motivacao da demanda": ", ".join(motivacoes),
    "Riscos": riscos,
    "Resultados": resultados,
    "Objetos e quantidades": objetos_quantidades,
}
pendentes = [nome for nome, valor in campos_minimos.items() if not str(valor).strip()]
if pendentes:
    with st.expander("Campos recomendados ainda nao preenchidos"):
        st.write("\n".join(f"- {campo}" for campo in pendentes))

if st.button("Gerar minuta de DOD", type="primary"):
    if not any(str(valor).strip() for valor in dados_dod.values()):
        st.warning("Preencha ao menos os dados principais da demanda antes de gerar a minuta.")
        st.stop()

    with st.spinner("Gerando minuta..."):
        try:
            client = obter_cliente_openai()
            trechos_institucionais = buscar_contexto_institucional(
                client,
                dados_dod,
                anexos_contexto,
            )
            contexto_institucional = formatar_contexto_institucional(trechos_institucionais)
            resposta = client.responses.create(
                model=obter_modelo(),
                input=montar_prompt(
                    dados_dod,
                    contexto_repo,
                    anexos_contexto,
                    contexto_institucional,
                ),
                max_output_tokens=5000,
            )
            texto_resposta = resposta.output_text
        except RateLimitError as erro:
            mostrar_erro_openai("Erro 429 retornado pela OpenAI API. Pode ser limite de uso, quota ou billing_not_active.", erro)
        except AuthenticationError as erro:
            mostrar_erro_openai("Erro de autenticacao. Verifique a OPENAI_API_KEY nos Secrets do Streamlit.", erro)
        except BadRequestError as erro:
            mostrar_erro_openai("Requisicao invalida. Verifique se o modelo configurado existe e esta disponivel para sua conta.", erro)
        except APIStatusError as erro:
            mostrar_erro_openai("Erro retornado pela OpenAI API.", erro)
        except Exception as erro:
            st.error("Erro inesperado ao chamar a OpenAI API.")
            st.write(f"**Timestamp:** `{obter_timestamp_brasilia()}`")
            st.write(f"**OpenAI Python SDK:** `{obter_versao_openai_sdk()}`")
            st.code(str(erro))
            st.stop()

    st.subheader("Minuta gerada")
    if trechos_institucionais:
        with st.expander("Contexto institucional recuperado"):
            for item in trechos_institucionais:
                st.write(
                    f"**{item.get('tipo_documental')}** - {item.get('arquivo')} "
                    f"(score {item.get('score')})"
                )
                st.caption(item.get("assunto") or "Sem assunto informado")
    else:
        st.info("Nenhum trecho da biblioteca institucional foi usado nesta geracao.")

    st.markdown(texto_resposta)
    st.download_button("Baixar minuta em Markdown", data=texto_resposta, file_name="minuta_dod.md", mime="text/markdown")
